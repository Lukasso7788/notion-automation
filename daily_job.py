import os
import json
import re
from datetime import datetime, timedelta
from zoneinfo import ZoneInfo

import requests
from dotenv import load_dotenv
from docx import Document

# === LOAD ENV ===
load_dotenv()

NOTION_API_KEY = os.getenv("NOTION_API_KEY")
TASKS_DB_ID = os.getenv("TASKS_DB_ID")
DAILY_LOG_DB_ID = os.getenv("DAILY_LOG_DB_ID")
STRATEGY_DB_ID = os.getenv("STRATEGY_DB_ID")
TIMEZONE = os.getenv("TIMEZONE", "Europe/Kyiv")

# AI (DeepSeek via OpenRouter)
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
OPENROUTER_BASE_URL = "https://openrouter.ai/api/v1"
MODEL_NAME = "deepseek/deepseek-chat"

# Telegram
TELEGRAM_BOT_TOKEN = os.getenv("TELEGRAM_BOT_TOKEN")
TELEGRAM_CHAT_ID = os.getenv("TELEGRAM_CHAT_ID")

# Discord (optional)
DISCORD_WEBHOOK_URL = os.getenv("DISCORD_WEBHOOK_URL")

# Advice file (exported Notion page "ЗАМЕТКИ И СТРАТЕГИИ")
ADVICE_FILE_PATH = os.getenv("ADVICE_FILE_PATH", "notes_strategies.txt")

# Notion headers
HEADERS = {
    "Authorization": f"Bearer {NOTION_API_KEY}",
    "Notion-Version": "2022-06-28",
    "Content-Type": "application/json",
}

# ---------------------------------------------------------
# DAILY RECURRING TASKS TEMPLATES
# ---------------------------------------------------------
DAILY_RECURRING_TASKS = [
    {
        "name": "Утренний ритуал — прочитать план, записать таски",
        "planned": 10,
        "type": "Admin",
    },
    {
        "name": "Практика программирования / курсы",
        "planned": 120,
        "type": "Learning",
    },
    {
        "name": "Физуха",
        "planned": 60,
        "type": "Gym",
    },
    {
        "name": "Немецкий — продлить стрик",
        "planned": 20,
        "type": "Learning",
    },
    {
        "name": "Вечерний ритуал — прочитать summary, записать таски и инфу",
        "planned": 10,
        "type": "Admin",
    },
]


# =========================================================
# 📅 DATE HELPERS
# =========================================================
def get_today():
    tz = ZoneInfo(TIMEZONE)
    return datetime.now(tz).date()


def get_yesterday():
    return get_today() - timedelta(days=1)


def get_tomorrow():
    return get_today() + timedelta(days=1)


def get_target_day_for_summary():
    """Всегда делаем summary за вчерашний день."""
    return get_yesterday()


# =========================================================
# 🔧 BASIC NOTION API
# =========================================================
def query_database(db_id, payload=None):
    url = f"https://api.notion.com/v1/databases/{db_id}/query"
    res = requests.post(url, headers=HEADERS, json=payload or {})
    res.raise_for_status()
    return res.json()


def update_page(page_id, payload):
    url = f"https://api.notion.com/v1/pages/{page_id}"
    res = requests.patch(url, headers=HEADERS, json=payload)
    res.raise_for_status()
    return res.json()


def create_page(db_id, properties, children=None):
    url = "https://api.notion.com/v1/pages"
    body = {"parent": {"database_id": db_id}, "properties": properties}
    if children:
        body["children"] = children
    res = requests.post(url, headers=HEADERS, json=body)
    res.raise_for_status()
    return res.json()


# =========================================================
# 🧩 SAFE HELPERS
# =========================================================
def safe_select_name(props, field_name):
    field = props.get(field_name) or {}
    select = field.get("select") or {}
    return select.get("name")


def safe_number(props, field_name):
    field = props.get(field_name) or {}
    num = field.get("number")
    return num or 0


def safe_checkbox(props, field_name):
    field = props.get(field_name) or {}
    cb = field.get("checkbox")
    return bool(cb)


def clean_text(txt: str) -> str:
    if txt is None:
        return ""
    txt = txt.replace("\r", " ").replace("\n\n\n", "\n")
    txt = re.sub(r"[\x00-\x1f\x80-\xff]", "", txt)
    return txt.strip()


# =========================================================
# 📌 TASKS
# =========================================================
def get_tasks_for_date(date):
    payload = {
        "filter": {
            "property": "Date",
            "date": {"equals": date.isoformat()},
        }
    }
    data = query_database(TASKS_DB_ID, payload)
    return data["results"]


def ensure_daily_recurring_tasks(target_day):
    """
    Создаем фиксированный набор задач на target_day, если их нет.
    """
    created = 0
    for t in DAILY_RECURRING_TASKS:
        name = t["name"]
        planned = t["planned"]
        ttype = t["type"]

        payload = {
            "filter": {
                "and": [
                    {
                        "property": "Date",
                        "date": {"equals": target_day.isoformat()},
                    },
                    {
                        "property": "Name",
                        "title": {"equals": name},
                    },
                ]
            }
        }

        data = query_database(TASKS_DB_ID, payload)
        if data["results"]:
            continue

        props = {
            "Name": {"title": [{"text": {"content": name}}]},
            "Date": {"date": {"start": target_day.isoformat()}},
            "Status": {"select": {"name": "Todo"}},
            "Type": {"select": {"name": ttype}},
            "Auto-roll?": {"checkbox": False},
            "Rollovers": {"number": 0},
            "Planned duration (min)": {"number": planned},
            "Actual duration (min)": {"number": 0},
        }

        create_page(TASKS_DB_ID, props)
        created += 1

    return created


# =========================================================
# 🔁 AUTO-ROLL (из целевого дня → завтра)
# =========================================================
def auto_roll_tasks(tasks, target_day):
    tomorrow = target_day + timedelta(days=1)
    rolled_count = 0

    for task in tasks:
        props = task.get("properties", {})
        status = safe_select_name(props, "Status")
        auto_flag = safe_checkbox(props, "Auto-roll?")

        if status == "Done" or not auto_flag:
            continue

        page_id = task["id"]
        current_roll = safe_number(props, "Rollovers")

        update_page(
            page_id,
            {
                "properties": {
                    "Date": {"date": {"start": tomorrow.isoformat()}},
                    "Rollovers": {"number": current_roll + 1},
                }
            },
        )
        rolled_count += 1

    return rolled_count


# =========================================================
# 📊 STATS
# =========================================================
def calculate_stats(tasks):
    total = len(tasks)
    done = 0
    planned = 0
    actual = 0
    deep = 0

    for t in tasks:
        p = t.get("properties", {})

        status = safe_select_name(p, "Status")
        if status == "Done":
            done += 1

        planned += safe_number(p, "Planned duration (min)")
        a = safe_number(p, "Actual duration (min)")
        actual += a

        if safe_select_name(p, "Type") == "Deep work":
            deep += a

    return {
        "total": total,
        "done": done,
        "planned_min": planned,
        "actual_min": actual,
        "deep_work_min": deep,
    }


# =========================================================
# 📚 STRATEGY SNAPSHOT
# =========================================================
def load_strategy_snapshot():
    if not STRATEGY_DB_ID:
        return "Нет данных стратегии (STRATEGY_DB_ID не задан)."

    try:
        data = query_database(STRATEGY_DB_ID, {})
    except Exception as e:
        return f"Не удалось загрузить стратегию: {e}"

    lines = []
    for page in data.get("results", []):
        props = page.get("properties", {})
        name_parts = props.get("Name", {}).get("title", [])
        name = name_parts[0]["plain_text"] if name_parts else "Без названия"

        status = safe_select_name(props, "Status") or "-"
        priority = safe_select_name(props, "Priority") or "-"
        horizon = safe_select_name(props, "Horizon") or "-"

        line = f"{name} [Status: {status}, Priority: {priority}, Horizon: {horizon}]"
        lines.append(line)

    if not lines:
        return "Стратегия не заполнена."
    return "\n".join(lines[:50])

# =========================================================
# 🧠 AI CLIENT
# =========================================================
def ai_client():
    import openai

    return openai.OpenAI(
        base_url=OPENROUTER_BASE_URL,
        api_key=OPENAI_API_KEY,
    )


# ---------------------------------------------------------
# 🧠 AI COMMENT FOR TASK
# ---------------------------------------------------------
def ai_comment_for_task(task):
    client = ai_client()

    props = task.get("properties", {})
    name_parts = props.get("Name", {}).get("title", [])
    name = name_parts[0]["plain_text"] if name_parts else "Без названия"

    task_type = safe_select_name(props, "Type") or "-"
    complexity = safe_number(props, "Complexity")
    rollovers = safe_number(props, "Rollovers")
    planned = safe_number(props, "Planned duration (min)")

    prompt = f"""
Ты — мой строгий, но адекватный продакт-наставник.

Задача: "{name}"
Тип: {task_type}
Сложность: {complexity}
Переносов: {rollovers}
Плановое время: {planned} мин

Дай один короткий комментарий (1–2 предложения), без markdown и эмодзи:
- как лучше выполнить
- что важно учесть
- если задача слишком большая — предложи упрощение.

Ответь ОДНИМ параграфом без переносов строк.
"""

    resp = client.chat.completions.create(
        model=MODEL_NAME,
        messages=[{"role": "user", "content": prompt}],
        max_tokens=120,
        temperature=0.3,
    )

    text = resp.choices[0].message.content
    return clean_text(text)


# ---------------------------------------------------------
# 📖 LOAD ADVICE TEXT FROM FILE
# ---------------------------------------------------------
def load_advice_lines():
    """
    Читаем файл заметок, режем на строки/фразы, фильтруем по длине.
    """
    if not ADVICE_FILE_PATH or not os.path.exists(ADVICE_FILE_PATH):
        return []

    with open(ADVICE_FILE_PATH, "r", encoding="utf-8", errors="ignore") as f:
        content = f.read()

    raw_lines = [clean_text(x) for x in content.split("\n")]
    lines = [x for x in raw_lines if 40 <= len(x) <= 300]
    return lines


def pick_daily_advice(lines):
    import random

    if not lines:
        return ""
    return random.choice(lines)


# ---------------------------------------------------------
# 🧠 AI SUMMARY + PLAN (JSON)
# ---------------------------------------------------------
def generate_ai_summary_and_plan(stats, target_day, strategy_snapshot):
    client = ai_client()

    prompt = f"""
Ты — мой персональный ИИ-коуч и стратег.

Вот статистика дня ({target_day}):
{json.dumps(stats, ensure_ascii=False, indent=2)}

Вот краткий срез моей стратегии (из отдельной базы Strategy):
{strategy_snapshot}

Твоя задача:
1) Кратко и чётко описать, как прошёл день.
2) Оценить, насколько день соответствует долгосрочной стратегии.
3) Сформировать конкретный план на завтра.

Формат ответа СТРОГО в JSON:

{{
  "summary": "Краткий разбор дня.",
  "strategy_alignment": "Связь дня со стратегией.",
  "plan_tomorrow": [
    "Пункт плана 1",
    "Пункт плана 2"
  ]
}}
"""

    resp = client.chat.completions.create(
        model=MODEL_NAME,
        messages=[{"role": "user", "content": prompt}],
        max_tokens=700,
        temperature=0.4,
    )

    raw = resp.choices[0].message.content.strip()

    try:
        data = json.loads(raw)
    except Exception:
        return clean_text(raw), "", []

    summary = clean_text(data.get("summary", ""))
    strategy_alignment = clean_text(data.get("strategy_alignment", ""))
    plan_list = data.get("plan_tomorrow", [])
    if not isinstance(plan_list, list):
        plan_list = []

    return summary, strategy_alignment, [clean_text(x) for x in plan_list]


# =========================================================
# 🧾 DAILY LOG
# =========================================================
def determine_status(stats):
    if stats["total"] == 0:
        return "On track"
    r = stats["done"] / stats["total"]
    if r >= 0.9:
        return "Ahead"
    elif r >= 0.6:
        return "On track"
    return "Behind"


def create_daily_log(
    stats,
    summary,
    strategy_alignment,
    plan_list,
    target_day,
    daily_advice,
):
    plan_text = "\n".join(f"- {p}" for p in plan_list) if plan_list else ""

    props = {
        "Name": {"title": [{"text": {"content": f"Day {target_day}"}}]},
        "Date": {"date": {"start": target_day.isoformat()}},
        "Status vs plan": {"select": {"name": determine_status(stats)}},
        "Total tasks": {"number": stats["total"]},
        "Done tasks": {"number": stats["done"]},
        "Planned min": {"number": stats["planned_min"]},
        "Actual min": {"number": stats["actual_min"]},
        "Deep work min": {"number": stats["deep_work_min"]},
        "AI plan for tomorrow": {
            "rich_text": [{"text": {"content": plan_text}}]
        },
        "Raw data (JSON)": {
            "rich_text": [
                {"text": {"content": json.dumps(stats, ensure_ascii=False)}}
            ]
        },
    }

    children = []

    if summary:
        children.append(
            {
                "object": "block",
                "type": "paragraph",
                "paragraph": {
                    "rich_text": [{"type": "text", "text": {"content": summary}}]
                },
            }
        )

    if strategy_alignment:
        children.extend(
            [
                {
                    "object": "block",
                    "type": "heading_3",
                    "heading_3": {
                        "rich_text": [
                            {"type": "text", "text": {"content": "Стратегия и день"}}
                        ]
                    },
                },
                {
                    "object": "block",
                    "type": "paragraph",
                    "paragraph": {
                        "rich_text": [
                            {
                                "type": "text",
                                "text": {"content": strategy_alignment},
                            }
                        ]
                    },
                },
            ]
        )

    if daily_advice:
        children.extend(
            [
                {
                    "object": "block",
                    "type": "heading_3",
                    "heading_3": {
                        "rich_text": [
                            {"type": "text", "text": {"content": "Совет дня"}}
                        ]
                    },
                },
                {
                    "object": "block",
                    "type": "paragraph",
                    "paragraph": {
                        "rich_text": [
                            {"type": "text", "text": {"content": daily_advice}}
                        ]
                    },
                },
            ]
        )

    return create_page(DAILY_LOG_DB_ID, props, children)


# =========================================================
# 📲 TELEGRAM / DISCORD
# =========================================================
def send_telegram_message(text: str):
    if not TELEGRAM_BOT_TOKEN or not TELEGRAM_CHAT_ID:
        print("Telegram not configured, skipping send_telegram_message")
        return

    url = f"https://api.telegram.org/bot{TELEGRAM_BOT_TOKEN}/sendMessage"
    payload = {
        "chat_id": TELEGRAM_CHAT_ID,
        "text": text,
        "parse_mode": "Markdown",
    }

    try:
        res = requests.post(url, json=payload, timeout=15)
        if not res.ok:
            print("Telegram sendMessage error:", res.text)
    except Exception as e:
        print("Telegram sendMessage exception:", e)


def send_telegram_document(file_path: str, caption: str | None = None):
    if not TELEGRAM_BOT_TOKEN or not TELEGRAM_CHAT_ID:
        print("Telegram not configured, skipping send_telegram_document")
        return

    url = f"https://api.telegram.org/bot{TELEGRAM_BOT_TOKEN}/sendDocument"

    try:
        with open(file_path, "rb") as f:
            files = {"document": f}
            data = {"chat_id": TELEGRAM_CHAT_ID}
            if caption:
                data["caption"] = caption
            res = requests.post(url, data=data, files=files, timeout=30)
        if not res.ok:
            print("Telegram sendDocument error:", res.text)
    except Exception as e:
        print("Telegram sendDocument exception:", e)


def _truncate_for_discord(content: str, limit: int = 2000) -> str:
    if content is None:
        return ""
    if len(content) <= limit:
        return content
    return content[: limit - 20] + "\n...(truncated)..."

# =========================================================
# 📄 DOCX GENERATION
# =========================================================
def build_plan_docx(
    plan_day,
    summary_day,
    tasks_for_day,
    plan_list,
    daily_advice,
):
    filename = f"plan_{plan_day.isoformat()}.docx"
    doc = Document()

    doc.add_heading(f"Plan for {plan_day}", level=1)

    # AI plan
    doc.add_heading("AI Plan", level=2)
    if plan_list:
        for item in plan_list:
            doc.add_paragraph(item, style="List Bullet")
    else:
        doc.add_paragraph("No explicit plan from AI.")

    # Tasks
    doc.add_heading("Tasks for the day", level=2)
    if not tasks_for_day:
        doc.add_paragraph("No tasks found.")
    else:
        for t in tasks_for_day:
            name = t["name"]
            ttype = t["type"]
            planned = t["planned"]
            comment = t["comment"]
            advice = t.get("advice", "")

            p = doc.add_paragraph(style="List Number")
            p.add_run(f"{name} [{ttype}] — {planned} min").bold = True

            if comment:
                doc.add_paragraph(f"AI comment: {comment}")
            if advice:
                doc.add_paragraph(f"Advice: {advice}")

    if daily_advice:
        doc.add_heading("Daily Advice", level=2)
        doc.add_paragraph(daily_advice)

    doc.save(filename)
    return filename


# =========================================================
# 🌅 TASKS FOR DAY: AI COMMENT + ADVICE
# =========================================================
def prepare_tasks_for_day(day, advice_lines):
    tasks = get_tasks_for_date(day)
    print(f"Found {len(tasks)} tasks for day ({day})")

    enriched = []

    for task in tasks:
        props = task.get("properties", {})
        name_parts = props.get("Name", {}).get("title", [])
        name = name_parts[0]["plain_text"] if name_parts else "Без названия"

        task_type = safe_select_name(props, "Type") or "-"
        planned = safe_number(props, "Planned duration (min)")

        # AI comment
        try:
            comment = ai_comment_for_task(task)
        except Exception as e:
            print(f"AI comment failed for task '{name}': {e}")
            comment = ""

        advice = ""

        # save AI comment to Notion
        try:
            update_page(
                task["id"],
                {
                    "properties": {
                        "AI comment": {
                            "rich_text": [
                                {"text": {"content": comment or ""}}
                            ]
                        }
                    }
                },
            )
        except Exception as e:
            print(f"Failed to update AI comment in Notion for '{name}': {e}")

        enriched.append(
            {
                "name": name,
                "type": task_type,
                "planned": planned,
                "comment": comment,
                "advice": advice,
            }
        )

    daily_advice = pick_daily_advice(advice_lines)
    return enriched, daily_advice


# =========================================================
# 🚀 MAIN
# =========================================================
def main():
    today = get_today()
    summary_day = get_target_day_for_summary()  # ВЧЕРА
    plan_day = today                             # СЕГОДНЯ

    print(
        f"\n=== RUNNING DAILY JOB "
        f"(today={today}, summary_for={summary_day}, plan_for={plan_day}) ===\n"
    )

    # 0) Load advice
    advice_lines = load_advice_lines()
    if advice_lines:
        print(f"Loaded {len(advice_lines)} advice lines")
    else:
        print("No advice lines loaded")

    # 1) Yesterday tasks → stats + autoroll
    tasks_yesterday = get_tasks_for_date(summary_day)
    print(f"Loaded {len(tasks_yesterday)} tasks for {summary_day}")

    rolled = auto_roll_tasks(tasks_yesterday, summary_day)
    print(f"Rolled over {rolled} tasks")

    stats = calculate_stats(tasks_yesterday)
    print("Stats:", stats)

    # 2) Strategy snapshot
    strategy_snapshot = load_strategy_snapshot()
    print("Strategy snapshot loaded")

    # 3) AI summary + PLAN НА СЕГОДНЯ
    summary, strategy_alignment, plan_list = generate_ai_summary_and_plan(
        stats, summary_day, strategy_snapshot
    )
    print("AI summary + plan generated")

    # 4) Daily log (за вчера)
    daily_advice_for_log = pick_daily_advice(advice_lines)
    daily_log_page = create_daily_log(
        stats,
        summary,
        strategy_alignment,
        plan_list,
        summary_day,
        daily_advice_for_log,
    )
    print("Daily log created:", daily_log_page.get("id"))

    # 5) Ensure recurring tasks НА СЕГОДНЯ (FIX)
    created_recurring = ensure_daily_recurring_tasks(plan_day)
    print(f"Created {created_recurring} recurring tasks for {plan_day}")

    # 6) Prepare tasks НА СЕГОДНЯ (FIX)
    tasks_today, daily_advice = prepare_tasks_for_day(plan_day, advice_lines)
    print(
        f"Prepared {len(tasks_today)} tasks for {plan_day}; "
        f"daily advice: {bool(daily_advice)}"
    )

    # 7) Send plan message
    if tasks_today:
        lines = [f"*План задач на {plan_day}:*"]
        for t in tasks_today:
            line = f"- *{t['name']}* [{t['type']}] — {t['planned']} мин"
            if t["comment"]:
                line += f"\n    _{t['comment']}_"
            lines.append(line)
        if daily_advice:
            lines.append(f"\n*Совет дня:* {daily_advice}")
        tasks_message = "\n".join(lines)
    else:
        tasks_message = f"На {plan_day} задач не найдено."
        if daily_advice:
            tasks_message += f"\n\nСовет дня: {daily_advice}"

    send_telegram_message(tasks_message)
    send_discord_message(tasks_message)

    # 8) DOCX
    docx_path = build_plan_docx(
        plan_day=plan_day,
        summary_day=summary_day,
        tasks_for_day=tasks_today,
        plan_list=plan_list,
        daily_advice=daily_advice,
    )
    print("DOCX generated:", docx_path)

    send_telegram_document(docx_path, caption=f"План на {plan_day}")
    send_discord_file(docx_path, content=f"План на {plan_day}")

    print("\n=== DONE ===\n")


if __name__ == "__main__":
    main()
